"""
Gera documentação de apresentação do Freky em estilo Samsung.
Samsung style: azul #1428A0, branco, preto, fonte Arial, design limpo e minimalista.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Cores Samsung ──────────────────────────────────────────────
SAMSUNG_BLUE   = RGBColor(0x14, 0x28, 0xA0)
SAMSUNG_LIGHT  = RGBColor(0x00, 0x6A, 0xC6)
DARK_BG        = RGBColor(0x0A, 0x0A, 0x0A)
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY     = RGBColor(0xF5, 0xF5, 0xF5)
MID_GRAY       = RGBColor(0xCC, 0xCC, 0xCC)
TEXT_DARK      = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT_CYAN    = RGBColor(0x00, 0xB9, 0xE4)

def rgb_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)

def set_cell_borders(cell, color="1428A0", size=6):  # noqa
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_table_borders(table, color="CCCCCC", size=4):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), str(size))
                border.set(qn('w:color'), color)
                tcBorders.append(border)
            tcPr.append(tcBorders)

def add_paragraph(doc, text="", bold=False, italic=False, size=11,
                  color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6, font_name="Arial"):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color
    return p

def add_heading(doc, text, level=1):
    sizes = {1: 26, 2: 18, 3: 13}
    colors = {1: SAMSUNG_BLUE, 2: SAMSUNG_BLUE, 3: TEXT_DARK}
    bolds = {1: True, 2: True, 3: True}
    before = {1: 18, 2: 14, 3: 10}
    after = {1: 8, 2: 6, 3: 4}
    return add_paragraph(doc, text,
                         bold=bolds.get(level, True),
                         size=sizes.get(level, 13),
                         color=colors.get(level, TEXT_DARK),
                         space_before=before.get(level, 10),
                         space_after=after.get(level, 4))

def add_rule(doc, color=SAMSUNG_BLUE, thickness=18):
    """Linha horizontal decorativa."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    hex_color = rgb_hex(color)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_table(doc, headers, rows, header_bg=SAMSUNG_BLUE, header_text=WHITE,
              stripe=True):
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # Cabeçalho
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(h)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = header_text

    # Linhas de dados
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = LIGHT_GRAY if (stripe and r_idx % 2 == 0) else WHITE
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(str(val))
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            run.font.color.rgb = TEXT_DARK

    set_table_borders(table, color="CCCCCC", size=4)

    # Ajusta larguras de colunas (distribui igualmente)
    table_width = Cm(16)
    col_width = table_width // col_count
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    return table

def add_code_block(doc, code_text):
    """Bloco de código estilo terminal Samsung (fundo escuro)."""
    lines = code_text.strip().split('\n')
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, RGBColor(0x1E, 0x1E, 0x2E))
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = ACCENT_CYAN
    set_cell_borders(cell, color="1428A0", size=8)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_info_box(doc, text, bg=RGBColor(0xE8, 0xEE, 0xF9)):
    """Caixa de destaque azul claro."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_cell_borders(cell, color="1428A0", size=12)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = SAMSUNG_BLUE
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_bullet(doc, items, color=SAMSUNG_BLUE):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("▸  ")
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = color
        run2 = p.add_run(item)
        run2.font.name = "Arial"
        run2.font.size = Pt(10)
        run2.font.color.rgb = TEXT_DARK

# ══════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════

doc = Document()

# Margens
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)

# ── CAPA ────────────────────────────────────────────────────────
# Banner azul de topo
cover_table = doc.add_table(rows=1, cols=1)
cover_table.style = 'Table Grid'
cover_cell = cover_table.rows[0].cells[0]
set_cell_bg(cover_cell, SAMSUNG_BLUE)
cover_cell.width = Cm(16)

p_logo = cover_cell.paragraphs[0]
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_logo.paragraph_format.space_before = Pt(30)
p_logo.paragraph_format.space_after = Pt(4)
run_logo = p_logo.add_run("FREKY")
run_logo.font.name = "Arial"
run_logo.font.size = Pt(52)
run_logo.font.bold = True
run_logo.font.color.rgb = WHITE

p_sub = cover_cell.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(4)
run_sub = p_sub.add_run("Sistema de Detecção X-Ray com IA")
run_sub.font.name = "Arial"
run_sub.font.size = Pt(16)
run_sub.font.color.rgb = ACCENT_CYAN

p_desc = cover_cell.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_desc.paragraph_format.space_before = Pt(0)
p_desc.paragraph_format.space_after = Pt(30)
run_desc = p_desc.add_run("HI-SCAN 6040i  ·  YOLOv8 ONNX  ·  FastAPI  ·  React")
run_desc.font.name = "Arial"
run_desc.font.size = Pt(11)
run_desc.font.color.rgb = RGBColor(0xBB, 0xCC, 0xEE)

doc.add_paragraph()

add_paragraph(doc, "Detecção automática de itens eletrônicos em bagagens",
              bold=True, size=14, color=TEXT_DARK,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
add_paragraph(doc, "Versão 1.0  ·  Março 2026",
              size=10, color=RGBColor(0x80, 0x80, 0x80),
              align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=24)

add_rule(doc, color=SAMSUNG_BLUE, thickness=24)
doc.add_page_break()

# ── 1. SOBRE O PROJETO ───────────────────────────────────────────
add_heading(doc, "1. Sobre o Projeto", 1)
add_rule(doc, color=SAMSUNG_BLUE)

add_paragraph(doc,
    "O Freky é um sistema de visão computacional em tempo real desenvolvido para "
    "automatizar a identificação de dispositivos eletrônicos em imagens X-Ray "
    "geradas pelo scanner de bagagem HI-SCAN 6040i (Smiths Detection).",
    size=11, color=TEXT_DARK, space_after=8)

add_paragraph(doc,
    "O sistema processa cada imagem em menos de 2 segundos, classifica o resultado "
    "em três decisões operacionais e transmite o resultado instantaneamente ao "
    "operador via WebSocket.",
    size=11, color=TEXT_DARK, space_after=12)

add_heading(doc, "Decisões do sistema", 2)

add_table(doc,
    headers=["Decisão", "Significado", "Ação recomendada"],
    rows=[
        ["LIBERADO",     "Nenhum item restrito detectado",     "Bagagem pode prosseguir"],
        ["VERIFICAR",    "Item restrito identificado",         "Inspeção manual obrigatória"],
        ["INCONCLUSIVO", "Confiança insuficiente",             "Revisão pelo operador"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Benefícios operacionais", 2)
add_bullet(doc, [
    "Redução do tempo médio de triagem",
    "Rastreabilidade completa com histórico de auditorias",
    "Feedback dos operadores para melhoria contínua do modelo",
    "Dashboard em tempo real com métricas operacionais",
    "Arquitetura containerizada — deploy em qualquer ambiente",
])

doc.add_page_break()

# ── 2. ARQUITETURA ───────────────────────────────────────────────
add_heading(doc, "2. Arquitetura do Sistema", 1)
add_rule(doc, color=SAMSUNG_BLUE)

add_info_box(doc,
    "O sistema é composto por 5 serviços principais orquestrados via Docker Compose, "
    "garantindo isolamento, escalabilidade e facilidade de deploy.")

add_table(doc,
    headers=["Serviço", "Tecnologia", "Função"],
    rows=[
        ["API",         "Python 3.11 · FastAPI · ONNX Runtime", "Inferência, REST, WebSocket"],
        ["Watcher",     "Python 3.11 · watchdog · httpx",       "Monitora pasta do scanner"],
        ["Dashboard",   "React 18 · Vite · Tailwind · Recharts", "Interface do operador"],
        ["SQL Server",  "Microsoft SQL Server 2022 Express",    "Persistência dos scans"],
        ["Redis",       "Redis 7-alpine",                       "Cache de estatísticas"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Fluxo de processamento", 2)
add_code_block(doc,
"""Scanner HI-SCAN 6040i
  │  deposita imagem em /scans/incoming/
  ▼
Watcher  →  aguarda estabilização  →  POST /scans/
  ▼
API  →  valida arquivo  →  ONNX YOLOv8  →  apply_rules()
  ▼
SQL Server (persiste)   Redis (invalida cache)   WebSocket (notifica)
  ▼
Dashboard do Operador  →  Decisão em tempo real""")

add_heading(doc, "Observabilidade (opcional)", 2)
add_bullet(doc, [
    "Prometheus :9090 — coleta métricas da API a cada 15s",
    "Grafana :3001 — dashboard provisionado automaticamente",
    "Redis Exporter :9121 — métricas de cache",
    "Health check /health/ready — verifica Banco + Redis + Modelo",
])

doc.add_page_break()

# ── 3. MODELO DE IA ──────────────────────────────────────────────
add_heading(doc, "3. Modelo de Inteligência Artificial", 1)
add_rule(doc, color=SAMSUNG_BLUE)

add_paragraph(doc,
    "O modelo de detecção é baseado na arquitetura YOLOv8, fine-tuned sobre o "
    "dataset HiXray contendo 8 classes de itens eletrônicos típicos em bagagens. "
    "O modelo é exportado para o formato ONNX para inferência otimizada.",
    size=11, color=TEXT_DARK, space_after=10)

add_heading(doc, "Classes detectadas", 2)
add_table(doc,
    headers=["ID", "Classe", "Categoria"],
    rows=[
        ["0", "portable_charger_1", "Restrito"],
        ["1", "portable_charger_2", "Restrito"],
        ["2", "mobile_phone",       "Restrito"],
        ["3", "laptop",             "Restrito"],
        ["4", "tablet",             "Restrito"],
        ["5", "cosmetic",           "Permitido"],
        ["6", "water",              "Permitido"],
        ["7", "nonmetallic_lighter","Monitorado"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Especificações técnicas", 2)
add_table(doc,
    headers=["Parâmetro", "Valor"],
    rows=[
        ["Arquitetura",         "YOLOv8m (fine-tuning)"],
        ["Tamanho de entrada",  "640 × 640 px"],
        ["Formato de saída",    "ONNX opset 17"],
        ["Threshold padrão",    "0.60 (global)"],
        ["Threshold alto",      "0.85 (itens críticos)"],
        ["Augmentação",         "Sem flip, sem cor — específico para X-ray"],
        ["Providers ONNX",      "CUDA → fallback CPU"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Pipeline de treinamento", 2)
add_bullet(doc, [
    "Dataset HiXray → convert_hixray_to_yolo.py (formato YOLO)",
    "validate_dataset.py (integridade das anotações)",
    "augment_xray.py (rotação ±10°, brilho ±30%)",
    "train.py — YOLOv8 fine-tuning, 50 épocas, early stopping patience=15",
    "evaluate.py — mAP@0.5, precision, recall, confusion matrix",
    "export_onnx.py — exporta para freky.onnx (opset=17, simplify=True)",
])

doc.add_page_break()

# ── 4. API ENDPOINTS ─────────────────────────────────────────────
add_heading(doc, "4. API — Endpoints Principais", 1)
add_rule(doc, color=SAMSUNG_BLUE)

add_heading(doc, "Autenticação e Segurança", 2)
add_table(doc,
    headers=["Método", "Rota", "Descrição"],
    rows=[
        ["POST", "/auth/login",              "Login JWT (rate limit: 10/min por IP)"],
        ["WS",   "/ws?token=<jwt>",          "Stream de resultados em tempo real"],
        ["GET",  "/health/ready",            "Readiness: Banco + Redis + Modelo"],
        ["GET",  "/metrics",                 "Métricas Prometheus"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Scans", 2)
add_table(doc,
    headers=["Método", "Rota", "Descrição"],
    rows=[
        ["POST", "/scans/",                  "Upload de imagem → retorna decisão"],
        ["POST", "/scans/{id}/feedback",     "Feedback do operador (confirmed/FP/FN)"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Auditoria", 2)
add_table(doc,
    headers=["Método", "Rota", "Descrição"],
    rows=[
        ["GET", "/audit/",              "Lista paginada (filtros: datas, decisão)"],
        ["GET", "/audit/stats",         "Totais por decisão (cache Redis 5min)"],
        ["GET", "/audit/daily?days=14", "Série temporal diária"],
        ["GET", "/audit/export",        "Download CSV completo com filtros"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Exemplo de resposta — POST /scans/", 2)
add_code_block(doc,
"""{
  "id": "3fa85f64-5717-4562-b3fc-2c963f66afa6",
  "filename": "scan_20260325_143021.jpg",
  "decision": "VERIFICAR",
  "processing_time_ms": 187.4,
  "detections": [
    {
      "class_name": "mobile_phone",
      "confidence": 0.92,
      "bbox": [145, 210, 312, 387]
    }
  ]
}""")

doc.add_page_break()

# ── 5. DASHBOARD ─────────────────────────────────────────────────
add_heading(doc, "5. Dashboard do Operador", 1)
add_rule(doc, color=SAMSUNG_BLUE)

add_paragraph(doc,
    "Interface React com atualização em tempo real via WebSocket. "
    "Autenticação JWT com controle de acesso por perfil (admin / operator).",
    size=11, color=TEXT_DARK, space_after=10)

add_heading(doc, "Páginas", 2)
add_table(doc,
    headers=["Página", "Função"],
    rows=[
        ["Login",              "Autenticação JWT, token persistido em localStorage"],
        ["Visão do Operador",  "Scan atual em tempo real + histórico dos últimos 50"],
        ["Auditoria",          "Histórico filtrado por data/decisão, paginação, feedback"],
        ["Estatísticas",       "Gráfico de barras diário + pizza por decisão + KPIs"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Funcionalidades em tempo real", 2)
add_bullet(doc, [
    "Conexão WebSocket com reconnect automático (backoff exponencial 1s → 30s)",
    "Indicador de status de conexão na barra de navegação",
    "Atualização instantânea do painel ao receber novo scan",
    "Histórico dos últimos 50 scans com modal de detalhes",
])

doc.add_paragraph()
add_heading(doc, "Feedback do Operador", 2)
add_table(doc,
    headers=["Tipo de Feedback", "Significado"],
    rows=[
        ["confirmed",      "Decisão correta — confirma o resultado do modelo"],
        ["false_positive",  "Item detectado erroneamente — não era restrito"],
        ["false_negative",  "Item restrito não detectado — modelo falhou"],
    ]
)

doc.add_page_break()

# ── 6. INFRAESTRUTURA ────────────────────────────────────────────
add_heading(doc, "6. Infraestrutura e Deploy", 1)
add_rule(doc, color=SAMSUNG_BLUE)

add_heading(doc, "Serviços Docker", 2)
add_table(doc,
    headers=["Serviço", "Porta", "Volume persistente"],
    rows=[
        ["API (FastAPI)",     "8000", "/scans, model/weights (ro)"],
        ["Watcher",          "—",    "/scans"],
        ["Dashboard (Nginx)", "3000", "—"],
        ["SQL Server 2022",  "1433", "sqlserver_data"],
        ["Redis 7",          "6379", "redis_data"],
        ["Prometheus",       "9090", "prometheus_data (30 dias)"],
        ["Grafana",          "3001", "grafana_data"],
    ]
)

doc.add_paragraph()
add_heading(doc, "Comandos de deploy", 2)
add_code_block(doc,
"""# 1. Clonar repositório e configurar ambiente
git clone https://github.com/GiovaneXavier/Freky.git
cd Freky
cp .env.example .env
# Editar .env com suas credenciais

# 2. Colocar modelo treinado
mkdir -p model/weights
cp /caminho/para/freky.onnx model/weights/freky.onnx

# 3. Iniciar todos os serviços
docker compose up -d

# 4. Verificar status
docker compose ps
curl http://localhost:8000/health/ready""")

doc.add_paragraph()
add_heading(doc, "Variantes de ambiente", 2)
add_table(doc,
    headers=["Arquivo", "Uso"],
    rows=[
        ["docker-compose.yml",         "Produção"],
        ["docker-compose.dev.yml",     "Desenvolvimento (hot-reload, portas extras)"],
        ["docker-compose.staging.yml", "Homologação"],
    ]
)

doc.add_paragraph()
add_heading(doc, "CI/CD — GitHub Actions", 2)
add_table(doc,
    headers=["Workflow", "Trigger", "Ações"],
    rows=[
        ["ci.yml",        "Push / PR → main",        "Lint (ruff) + Testes API + Testes Watcher"],
        ["docker.yml",    "Push main / tag v*.*.*",   "Build imagens + Push ghcr.io + Trivy scan"],
        ["model-eval.yml","Manual / Seg 06:00 UTC",   "Avaliação mAP, artefatos de métricas"],
    ]
)

doc.add_page_break()

# ── 7. SEGURANÇA ─────────────────────────────────────────────────
add_heading(doc, "7. Segurança", 1)
add_rule(doc, color=SAMSUNG_BLUE)

add_table(doc,
    headers=["Camada", "Controle implementado"],
    rows=[
        ["Autenticação",    "JWT HS256, expiração 8h, segredo mínimo 32 chars"],
        ["Rate limiting",   "10 req/min por IP no endpoint de login (slowapi)"],
        ["Senhas",          "bcrypt em produção, texto puro bloqueado (FREKY_ENV != test)"],
        ["CORS",            "Lista configurável de origens via ALLOWED_ORIGINS"],
        ["Upload",          "Limite 50 MB + PIL.Image.verify() antes de persistir"],
        ["Container scan",  "Trivy CRITICAL/HIGH — SARIF publicado no GitHub Security"],
        ["Banco de dados",  "SQL Server com credenciais via variável de ambiente"],
    ]
)

doc.add_paragraph()
add_info_box(doc,
    "Em produção, o JWT_SECRET_KEY é validado na inicialização da API. "
    "Valores padrão ou com menos de 32 caracteres causam erro imediato ao subir o serviço.")

doc.add_page_break()

# ── 8. PRÓXIMOS PASSOS ───────────────────────────────────────────
add_heading(doc, "8. Próximos Passos", 1)
add_rule(doc, color=SAMSUNG_BLUE)

add_heading(doc, "Modelo", 2)
add_bullet(doc, [
    "Treinar com dataset HiXray real (imagens do HI-SCAN 6040i)",
    "Ajustar thresholds por classe após avaliação em produção",
    "Publicar pesos como release no GitHub para CI/CD de avaliação",
])

doc.add_paragraph()
add_heading(doc, "Infraestrutura", 2)
add_bullet(doc, [
    "Configurar deploy em ambiente de homologação (docker-compose.staging.yml)",
    "Adicionar monitoramento de alertas no Grafana (regras de alerta)",
    "Implementar backup automático do SQL Server",
    "Configurar TLS/HTTPS no Nginx do dashboard",
])

doc.add_paragraph()
add_heading(doc, "Produto", 2)
add_bullet(doc, [
    "Validar pipeline de inferência com imagens reais do scanner",
    "Coletar feedback dos operadores para ciclo de retreinamento",
    "Adicionar exportação PDF de relatórios de auditoria",
    "Internacionalização do dashboard (i18n)",
])

doc.add_page_break()

# ── 9. RESUMO TÉCNICO ────────────────────────────────────────────
add_heading(doc, "9. Resumo das Tecnologias", 1)
add_rule(doc, color=SAMSUNG_BLUE)

add_table(doc,
    headers=["Camada", "Tecnologia", "Versão"],
    rows=[
        ["Backend",        "Python / FastAPI / Uvicorn",          "3.11 / 0.115 / 0.30"],
        ["Inferência",     "ONNX Runtime / YOLOv8 (ultralytics)", "1.19 / 8.3"],
        ["Banco de dados", "SQL Server / SQLAlchemy async",       "2022 / 2.0"],
        ["Cache",          "Redis",                               "7-alpine"],
        ["Frontend",       "React / TypeScript / Vite / Tailwind","18 / 5 / 5 / 3"],
        ["Gráficos",       "Recharts",                            "2.x"],
        ["Auth",           "python-jose / passlib bcrypt",        "3.3 / 1.7"],
        ["Rate limit",     "slowapi",                             "0.1.9"],
        ["Métricas",       "Prometheus / Grafana",                "2.54 / 11.2"],
        ["Containers",     "Docker / Docker Compose",             "—"],
        ["CI/CD",          "GitHub Actions",                      "—"],
        ["Lint/Test",      "ruff / pytest / pytest-cov",          "0.6 / 8.3 / 5.0"],
    ]
)

doc.add_paragraph()

# Rodapé final
add_rule(doc, color=SAMSUNG_BLUE, thickness=12)
add_paragraph(doc,
    "Freky  ·  Sistema de Detecção X-Ray com IA  ·  Versão 1.0  ·  Março 2026",
    size=9, color=RGBColor(0x80, 0x80, 0x80),
    align=WD_ALIGN_PARAGRAPH.CENTER)

# ── SALVA ────────────────────────────────────────────────────────
output_path = "/home/user/Freky/docs/Freky_Apresentacao_Samsung.docx"
doc.save(output_path)
print(f"Documento gerado: {output_path}")
